import pdfplumber
from docx import Document
from pathlib import Path


class ResumeParser:
    """
    Handles resume parsing for PDF and DOCX files.
    """

    def __init__(self):
        pass

    def parse_resume(self, file_path):
        """
        Detect file type and parse the resume.
        """

        file_path = Path(file_path)

        if not file_path.exists():
            return {
                "success": False,
                "message": "File not found."
            }

        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self._parse_pdf(file_path)

        elif extension == ".docx":
            return self._parse_docx(file_path)

        else:
            return {
                "success": False,
                "message": "Unsupported file type."
            }

    def _parse_pdf(self, file_path):

        text = ""

        try:

            with pdfplumber.open(file_path) as pdf:

                for page in pdf.pages:

                    page_text = page.extract_text()

                    if page_text:
                        text += page_text + "\n"

            return {
                "success": True,
                "file_name": file_path.name,
                "file_type": "PDF",
                "pages": len(pdf.pages),
                "text": text
            }

        except Exception as e:

            return {
                "success": False,
                "message": str(e)
            }

    def _parse_docx(self, file_path):

        try:

            document = Document(file_path)

            paragraphs = []

            for para in document.paragraphs:
                paragraphs.append(para.text)

            text = "\n".join(paragraphs)

            return {

                "success": True,
                "file_name": file_path.name,
                "file_type": "DOCX",
                "pages": len(document.paragraphs),
                "text": text

            }

        except Exception as e:

            return {

                "success": False,
                "message": str(e)

            }